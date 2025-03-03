import streamlit as st
import datetime
import re
import io
from docx import Document
from docx.shared import Pt
from components.unclassed import display_banner
import os
from dotenv import load_dotenv

load_dotenv()

display_banner()

template_file = "assets/TEMPLATE.docx"

try:
    from openai import OpenAI
except ImportError:
    st.error("The openai library is not installed. Install it with pip install openai")
    st.stop()

# Constants
API_KEY = os.getenv("api_key")

SYSTEM_PROMPT = """
Create a clear and simple Standard Operating Procedure (SOP)

Each SOP should have the following sections:

1. Procedures
   - Brief overview, then numbered steps.

2. Definitions (optional)
   - Explain unclear terms or acronyms.

Note: These will be US Air Force specific SOPs. Be sure to search for all relevant information from the Air Force from AFIs and DAFMANs.


- Search the user provided title and actions and provide the most accurate and up to date information based on their starting point, if given.
- Return only the SOP content, no other text.
- Relentlessly optimize for clarity and simplicity.
- Format: Plain text, no citations, no markdown.
- Do not include information that can become outdated.

"""

def generate_sop(title: str, actions: str) -> str:
    """Generate SOP suggestion by calling the API.

    Args:
        title (str): Title of the SOP.
        actions (str): Starting actions provided by the user.

    Returns:
        str: SOP suggestion with <think> parts removed.
    """
    messages = [
        {"role": "system", "content": SYSTEM_PROMPT},
        {"role": "user", "content": f"User has provided the following title and actions as a starting point for the SOP: {title} - {actions}"}
    ]
    client = OpenAI(api_key=API_KEY, base_url="https://api.perplexity.ai")
    response = client.chat.completions.create(
        model="sonar-reasoning-pro",
        messages=messages,
    )
    sop_text = response.choices[0].message.content
    sop_text = re.sub(r"<think>.*?</think>", "", sop_text, flags=re.DOTALL).strip()
    return sop_text


def replace_placeholders(text: str, field_values: dict) -> str:
    """Replace placeholders in the given text with field values.

    Args:
        text (str): Original text containing placeholders.
        field_values (dict): Mapping of field keys to values.

    Returns:
        str: Updated text.
    """
    for key, value in field_values.items():
        placeholder = f"{{{{{key}}}}}"
        if value.strip():
            text = text.replace(placeholder, value)
        else:
            text = text.replace(placeholder, "")
    return text


def process_document(doc: Document, field_values: dict) -> None:
    """Process the docx Document, replacing placeholders with values.

    Args:
        doc (Document): The docx document to process.
        field_values (dict): Mapping of field keys to values.
    """
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        if any(f"{{{{{key}}}}}" in paragraph.text for key in field_values):
            paragraph.text = replace_placeholders(paragraph.text, field_values)
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.text = replace_placeholders(paragraph.text, field_values)


def adjust_font(doc, size=Pt(8)):
    """Adjust all runs in the document (including those in tables) to the specified font size.

    Args:
        doc (Document): The document to adjust.
        size (Pt): The desired font size.
    """
    # Adjust paragraphs outside of tables
    for para in doc.paragraphs:
        for run in para.runs:
            run.font.size = size
    # Adjust paragraphs within tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = size


# Session state defaults
if 'title' not in st.session_state:
    st.session_state['title'] = ""
if 'code' not in st.session_state:
    st.session_state['code'] = ""
if 'checklist_no' not in st.session_state:
    st.session_state['checklist_no'] = ""
if 'revision' not in st.session_state:
    st.session_state['revision'] = ""
if 'date' not in st.session_state:
    st.session_state['date'] = datetime.datetime.now().strftime("%d %B %Y").upper()
if 'position' not in st.session_state:
    st.session_state['position'] = ""
if 'actions_text' not in st.session_state:
    st.session_state['actions_text'] = ""
if 'sop_suggestion' not in st.session_state:
    st.session_state['sop_suggestion'] = ""
if 'expanded' not in st.session_state:
    st.session_state['expanded'] = True

st.title("SOP Creator Tool")

st.session_state['title'] = st.text_input("Title", st.session_state['title'], placeholder="Enter a title for this SOP")
st.session_state['code'] = st.text_input("Code", st.session_state['code'])
st.session_state['checklist_no'] = st.text_input("Checklist No.", st.session_state['checklist_no'], placeholder="100.001")
st.session_state['revision'] = st.text_input("Revision", st.session_state['revision'])
st.session_state['date'] = st.text_input("Creation Date", st.session_state['date'])
st.session_state['position'] = st.text_input("Position", st.session_state['position'])
st.session_state['actions_text'] = st.text_area("Actions:", st.session_state['actions_text'], height=200, placeholder="Steps needed to complete this SOP. Provide as much detail as possible before asking for AI help.")

# Generate SOP functionality remains
if st.button("Get Help From AI" ):
    if st.session_state['title'].strip() == "":
        st.error("Please enter a Title first.")
    else:
        with st.spinner('Loading'):
            try:
                suggestion = generate_sop(st.session_state['title'], st.session_state['actions_text'])
                st.session_state['sop_suggestion'] = suggestion
            except Exception as e:
                st.error(f"API call failed: {e}")
                st.stop()
if st.session_state["sop_suggestion"]:
    with st.expander("SOP Suggestion", expanded=st.session_state["expanded"]):
        st.code(st.session_state["sop_suggestion"], language="text")
        col1, col2 = st.columns(2)
        if col1.button("Accept"):
            st.session_state["actions_text"] = st.session_state["sop_suggestion"]
            st.toast("SOP updated from suggestion.")
            st.session_state["expanded"] = False
            st.rerun()
        if col2.button("Decline"):
            st.toast("SOP suggestion declined.")
            st.session_state["expanded"] = False

st.markdown("---")

def generate_doc():
    """Generate document from template and update session state with generated file."""
    required = ["title", "checklist_no", "date", "actions_text"]
    missing = [field for field in required if st.session_state[field].strip() == ""]
    if missing:
        st.error(f"Please fill in required fields: {', '.join(missing)}")
        return
    try:
        # Validate date format
        datetime.datetime.strptime(st.session_state['date'], "%d %B %Y")
    except ValueError:
        st.error("Please enter the date in DD MONTH YYYY format (e.g., 02 MARCH 2025)")
        return
    try:
        # Load the template document from the file
        with open(template_file, "rb") as f:
            doc = Document(io.BytesIO(f.read()))
        
        field_values = {
            "TITLE": st.session_state['title'],
            "CODE": st.session_state['code'],
            "CHECKLIST.NO": st.session_state['checklist_no'],
            "REV": st.session_state['revision'],
            "DATE": st.session_state['date'],
            "POSITION": st.session_state['position'],
            "ACTIONS": st.session_state['actions_text'],
        }
        process_document(doc, field_values)
        adjust_font(doc, Pt(8))
        filename = f"{st.session_state['checklist_no']} - {st.session_state['title']}.docx"
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        st.session_state['doc_bytes'] = buf.read()
        st.session_state['doc_filename'] = filename
    except Exception as e:
        st.error(f"Error generating document: {e}")

if st.button("Generate SOP"):
    generate_doc()
    if 'doc_bytes' in st.session_state and st.session_state['doc_bytes']:
        st.download_button(
            label="Download",
            data=st.session_state['doc_bytes'],
            file_name=st.session_state['doc_filename'],
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="download_button"
        )
    else:
        st.error("Document generation failed or missing required fields.")

with st.sidebar:
    st.markdown("Use this tool to speed up your SOP creation process. You can use the AI to help you create the SOP, and then download the final document.")

    st.markdown("---")

    st.markdown("""
Tip:
Provide a clear, descriptive title that outlines the specific procedure. In the actions section, include detailed steps and note any unique, organization-specific practices. This extra detail helps the AI generate an accurate and tailored SOP.

Example:
- Title: "Update CISCO 9300 Series Switch Firmware"
- Actions: "Download the latest firmware from {Airforce Approved URL} and update via USB."

- Outside of the Airforce, using TFTP, a web interface, etc. Would be the norm.
""")


